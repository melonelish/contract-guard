"""
批量将 contract_source.md 转换为 contract_source.docx (及可选 .pdf)

用法:
    python scripts/md_to_docx.py [--pdf]

依赖:
    pip install python-docx fpdf2
"""

import argparse
import os
import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF


def parse_md_to_docx(md_path: str, docx_path: str) -> bool:
    """将 Markdown 文件解析并写入 DOCX，保留标题/段落/列表/引用块结构。"""
    try:
        with open(md_path, "r", encoding="utf-8") as f:
            lines = f.readlines()
    except Exception as e:
        print(f"  [ERROR] 读取失败: {e}")
        return False

    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "SimSun"
    font.size = Pt(11)

    in_code_block = False
    code_buffer = []
    in_table = False
    table_rows = []

    def flush_code_block():
        nonlocal code_buffer
        if code_buffer:
            for line in code_buffer:
                p = doc.add_paragraph()
                run = p.add_run(line.rstrip())
                run.font.name = "Courier New"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            code_buffer = []

    def flush_table():
        nonlocal table_rows, in_table
        if table_rows and len(table_rows) >= 2:
            # Parse markdown table
            data_rows = []
            for row_str in table_rows:
                cells = [c.strip() for c in row_str.strip("|").split("|")]
                # Skip separator row (----)
                if all(re.match(r"^[-:]+$", c) for c in cells):
                    continue
                data_rows.append(cells)

            if data_rows:
                ncols = max(len(r) for r in data_rows)
                table = doc.add_table(rows=len(data_rows), cols=ncols)
                table.style = "Table Grid"
                for i, row_cells in enumerate(data_rows):
                    for j, cell_text in enumerate(row_cells):
                        if j < ncols:
                            table.rows[i].cells[j].text = cell_text
        table_rows = []
        in_table = False

    for line in lines:
        stripped = line.rstrip("\n")

        # Code block toggle
        if stripped.startswith("```"):
            if in_code_block:
                flush_code_block()
                in_code_block = False
            else:
                if in_table:
                    flush_table()
                in_code_block = True
            continue

        if in_code_block:
            code_buffer.append(stripped)
            continue

        # Table detection
        if "|" in stripped and stripped.strip().startswith("|"):
            if not in_table:
                in_table = True
            table_rows.append(stripped)
            continue
        else:
            if in_table:
                flush_table()

        # Empty line
        if not stripped.strip():
            continue

        # Headings
        heading_match = re.match(r"^(#{1,6})\s+(.+)", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            # Remove markdown bold/italic markers
            text = re.sub(r"\*{1,2}(.+?)\*{1,2}", r"\1", text)
            heading = doc.add_heading(text, level=min(level, 4))
            continue

        # Blockquote
        if stripped.startswith(">"):
            text = stripped.lstrip("> ").strip()
            text = re.sub(r"\*{1,2}(.+?)\*{1,2}", r"\1", text)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            continue

        # Horizontal rule
        if re.match(r"^[-*_]{3,}\s*$", stripped):
            doc.add_paragraph("─" * 50)
            continue

        # Unordered list
        list_match = re.match(r"^(\s*)[-*+]\s+(.+)", stripped)
        if list_match:
            indent = len(list_match.group(1))
            text = list_match.group(2).strip()
            text = re.sub(r"\*{1,2}(.+?)\*{1,2}", r"\1", text)
            p = doc.add_paragraph(style="List Bullet")
            if indent > 0:
                p.paragraph_format.left_indent = Inches(0.25 + indent * 0.15)
            p.clear()
            run = p.add_run(text)
            continue

        # Ordered list
        ol_match = re.match(r"^(\s*)\d+\.\s+(.+)", stripped)
        if ol_match:
            text = ol_match.group(2).strip()
            text = re.sub(r"\*{1,2}(.+?)\*{1,2}", r"\1", text)
            p = doc.add_paragraph(style="List Number")
            p.clear()
            run = p.add_run(text)
            continue

        # Normal paragraph
        text = stripped.strip()
        # Clean markdown formatting
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)  # bold
        text = re.sub(r"\*(.+?)\*", r"\1", text)  # italic
        text = re.sub(r"`(.+?)`", r"\1", text)  # inline code
        text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)  # links

        if text:
            doc.add_paragraph(text)

    # Flush remaining
    if in_code_block:
        flush_code_block()
    if in_table:
        flush_table()

    try:
        doc.save(docx_path)
        return True
    except Exception as e:
        print(f"  [ERROR] 保存DOCX失败: {e}")
        return False


def parse_md_to_pdf(md_path: str, pdf_path: str) -> bool:
    """使用 fpdf2 将 Markdown 直接转换为 PDF，支持中文。"""
    try:
        with open(md_path, "r", encoding="utf-8") as f:
            lines = f.readlines()
    except Exception as e:
        print(f"  [ERROR] 读取失败: {e}")
        return False

    # 查找系统中文字体
    font_paths = [
        "C:/Windows/Fonts/simsun.ttc",
        "C:/Windows/Fonts/simhei.ttf",
        "C:/Windows/Fonts/msyh.ttc",
        "/usr/share/fonts/truetype/noto/NotoSansCJK-Regular.ttc",
    ]
    font_path = None
    for fp in font_paths:
        if os.path.exists(fp):
            font_path = fp
            break
    if not font_path:
        print("  [WARN] 未找到中文字体，跳过PDF")
        return False

    try:
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=20)
        pdf.add_font("zh", "", font_path)
        pdf.add_font("zh", "B", font_path)
        pdf.add_page()

        in_code_block = False
        in_table = False
        table_lines = []

        def write_text(text, style="", size=11):
            pdf.set_font("zh", style, size)
            pdf.multi_cell(0, 6, text, new_x="LMARGIN", new_y="NEXT")

        def flush_table():
            nonlocal table_lines, in_table
            if not table_lines:
                in_table = False
                return

            # Parse rows
            data_rows = []
            for row_str in table_lines:
                cells = [c.strip() for c in row_str.strip("|").split("|")]
                if all(re.match(r"^[-:]+$", c) for c in cells):
                    continue
                data_rows.append(cells)
            table_lines = []
            in_table = False

            if not data_rows:
                return

            ncols = max(len(r) for r in data_rows)
            col_w = (pdf.w - pdf.l_margin - pdf.r_margin) / ncols

            for row in data_rows:
                max_h = 6
                # Calculate max height needed
                for cell in row:
                    pdf.set_font("zh", "", 9)
                    nlines = max(1, len(cell) * pdf.get_string_width("中") / col_w + 1)
                    max_h = max(max_h, int(nlines) * 5)

                for j, cell in enumerate(row):
                    if j < ncols:
                        x = pdf.l_margin + j * col_w
                        y = pdf.get_y()
                        pdf.set_font("zh", "", 9)
                        # Draw cell border
                        pdf.rect(x, y, col_w, max_h)
                        pdf.set_xy(x + 1, y + 1)
                        pdf.multi_cell(col_w - 2, 5, cell, new_x="LEFT", new_y="TOP")
                pdf.set_y(pdf.get_y() + max_h)

        for line in lines:
            stripped = line.rstrip("\n")

            # Code block toggle
            if stripped.startswith("```"):
                if in_code_block:
                    in_code_block = False
                else:
                    if in_table:
                        flush_table()
                    in_code_block = True
                continue

            if in_code_block:
                pdf.set_font("zh", "", 9)
                pdf.multi_cell(0, 5, stripped, new_x="LMARGIN", new_y="NEXT")
                continue

            # Table detection
            if "|" in stripped and stripped.strip().startswith("|"):
                if not in_table:
                    in_table = True
                table_lines.append(stripped)
                continue
            else:
                if in_table:
                    flush_table()

            # Empty line
            if not stripped.strip():
                pdf.ln(3)
                continue

            # Headings
            heading_match = re.match(r"^(#{1,6})\s+(.+)", stripped)
            if heading_match:
                level = len(heading_match.group(1))
                text = heading_match.group(2).strip()
                text = re.sub(r"\*{1,2}(.+?)\*{1,2}", r"\1", text)
                sizes = {1: 18, 2: 16, 3: 14, 4: 13, 5: 12, 6: 11}
                pdf.ln(4)
                write_text(text, style="B", size=sizes.get(level, 11))
                pdf.ln(2)
                continue

            # Blockquote
            if stripped.startswith(">"):
                text = stripped.lstrip("> ").strip()
                text = re.sub(r"\*{1,2}(.+?)\*{1,2}", r"\1", text)
                x = pdf.get_x()
                pdf.set_x(x + 10)
                pdf.set_font("zh", "", 10)
                pdf.multi_cell(0, 5, text, new_x="LMARGIN", new_y="NEXT")
                continue

            # Horizontal rule
            if re.match(r"^[-*_]{3,}\s*$", stripped):
                pdf.ln(2)
                pdf.line(pdf.l_margin, pdf.get_y(), pdf.w - pdf.r_margin, pdf.get_y())
                pdf.ln(2)
                continue

            # Unordered list
            list_match = re.match(r"^(\s*)[-*+]\s+(.+)", stripped)
            if list_match:
                indent = len(list_match.group(1))
                text = list_match.group(2).strip()
                text = re.sub(r"\*{1,2}(.+?)\*{1,2}", r"\1", text)
                pdf.set_x(pdf.l_margin + 5 + indent * 3)
                write_text(f"• {text}", size=11)
                continue

            # Ordered list
            ol_match = re.match(r"^(\s*)\d+\.\s+(.+)", stripped)
            if ol_match:
                text = ol_match.group(2).strip()
                text = re.sub(r"\*{1,2}(.+?)\*{1,2}", r"\1", text)
                pdf.set_x(pdf.l_margin + 5)
                num = re.match(r"(\d+)\.", stripped).group(1)
                write_text(f"{num}. {text}", size=11)
                continue

            # Normal paragraph
            text = stripped.strip()
            text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
            text = re.sub(r"\*(.+?)\*", r"\1", text)
            text = re.sub(r"`(.+?)`", r"\1", text)
            text = re.sub(r"\[(.+?)\]\(.+?\)", r"\1", text)

            if text:
                write_text(text, size=11)

        # Flush remaining
        if in_table:
            flush_table()

        pdf.output(pdf_path)
        return True
    except Exception as e:
        print(f"  [ERROR] PDF生成失败: {e}")
        return False


def main():
    parser = argparse.ArgumentParser(description="批量转换 contract_source.md → docx/pdf")
    parser.add_argument("--pdf", action="store_true", help="同时生成 PDF")
    parser.add_argument(
        "--root",
        default="D:/FaLvXM/07-testing/generated",
        help="扫描根目录",
    )
    args = parser.parse_args()

    root = Path(args.root)
    if not root.exists():
        print(f"目录不存在: {root}")
        sys.exit(1)

    # 扫描所有包含 contract_source.md 的目录
    md_files = sorted(root.rglob("contract_source.md"))
    print(f"找到 {len(md_files)} 份合同源文件\n")

    success_docx = 0
    success_pdf = 0
    failed = []
    results = []

    for md_path in md_files:
        contract_dir = md_path.parent
        contract_name = contract_dir.name
        docx_path = contract_dir / "contract_source.docx"
        pdf_path = contract_dir / "contract_source.pdf"

        print(f"[{contract_name}]")

        # 跳过已存在的 docx
        if docx_path.exists():
            print(f"  [SKIP] DOCX 已存在")
            success_docx += 1
        else:
            ok = parse_md_to_docx(str(md_path), str(docx_path))
            if ok:
                size_kb = docx_path.stat().st_size / 1024
                print(f"  [OK] DOCX ({size_kb:.1f} KB)")
                success_docx += 1
            else:
                failed.append((contract_name, "docx"))
                print(f"  [FAIL] DOCX 转换失败")

        # PDF
        if args.pdf:
            if pdf_path.exists():
                print(f"  [SKIP] PDF 已存在")
                success_pdf += 1
            else:
                ok = parse_md_to_pdf(str(md_path), str(pdf_path))
                if ok:
                    size_kb = pdf_path.stat().st_size / 1024
                    print(f"  [OK] PDF ({size_kb:.1f} KB)")
                    success_pdf += 1
                else:
                    failed.append((contract_name, "pdf"))

        results.append(contract_name)

    # 汇总
    print("\n" + "=" * 60)
    print(f"转换完成:")
    print(f"  DOCX 成功: {success_docx}/{len(md_files)}")
    if args.pdf:
        print(f"  PDF  成功: {success_pdf}/{len(md_files)}")
    if failed:
        print(f"  失败: {len(failed)}")
        for name, fmt in failed:
            print(f"    - {name} ({fmt})")
    print(f"\n脚本路径: {Path(__file__).resolve()}")


if __name__ == "__main__":
    main()
